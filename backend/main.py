import os
import pprint
from bson.objectid import ObjectId

from docx import Document
from pymongo import MongoClient
import pymongo as pm
import datetime
import django

"""
responses:
:errors 1000: can't register user
:errors 2000: no account matching login
:errors 2001: wrong password
:errors 2002: account not verified
:errors 2003: no account matching id
:errors 2004: account has not enough rights
:errors 3000: pieces already taken
:errors 3001: not a single piece
:errors 4040: a error
:success OK: no error occurred during operation
"""

BOOL_TO_ABB = ["ENG", "GER", "FRE", "ESP", "ITA", "JAP", "CHI"]


def register(name, surname, mi, email, langs, login, password, status, vk=None, tg=None, fb=None):
    """
    :param name: obvious
    :param surname: obvious
    :param mi: middle initials
    :param email: obvious
    :param langs: user languages
    :param login: obvious
    :param password: obvious
    :param status: status translator/chief/both
    :return: user id or None if failed
    """

    a = {"name": name,
         "surname": surname,
         "mi": mi,
         "email": email,
         "langs": langs,
         "login": login,
         "password": password,
         "status": status,
         "vk": vk,
         "tg": tg,
         "fb": fb,
         "translated": 0,
         "pieces": [],
         "verified": False}

    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    try:
        user_id = acc.insert_one(a).inserted_id
        return ("OK", user_id)
    except pm.errors.DuplicateKeyError:
        return ("1000",)


def log_in(login, password):
    """
    :param login:
    :param password:
    :return: user id or None if failed
    """
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    user = acc.find_one({"login": login})
    if not (user is None):
        if user["password"] == password and user["verified"]:
            if user["verified"]:
                return user["_id"]
            else:
                return ("2002",)
        else:
            return ("2001",)
    else:
        return ("2000",)


def verify(user_id, key, decision="ADMITTED"):
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    if key == "NICE":
        if decision == "ADMITTED":
            acc.update_one(
                {"_id": ObjectId(user_id)},
                {"$set": {"verified": True}})
        else:
            acc.delete_one({"_id": ObjectId(user_id)})
        return ("OK",)
    else:
        return ("2004",)


def verify_file(doc_id, user_id, file_data):
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    acc = db.accounts
    user = acc.find_one({"_id": ObjectId(user_id)})
    if is_there_any_body(user_id):
        if user["status"] == "chief":
            lang_storage.update_one({"_id": ObjectId(doc_id)}, {"$set": {"status": "TRANSLATED", "chief": user_id}})
            file = lang_storage.find_one({"_id": ObjectId(doc_id)})
            push_to_file_storage(file["path"], file_data)
            return ("OK",)
        else:
            return ("2004",)
    else:
        return ("2003",)


def is_there_any_body(uid):
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    return not(acc.find_one({"_id": uid}) is None)


def split_to_pieces(number, name, lang, doc):
    """
       :param number: id number of document
       :param name: file name
       :param lang: language one of ENG, RUS, ESP, JAP, etc.
       :param doc: document docx
       :return: pieces ids
   """

    ids = list()
    for i in range(len(doc.paragraphs)):
        did = push_to_db(number, name, "WAITING_PIECE", lang, txt=doc.paragraphs[i].text, index=i, freedom=True)
        ids.append(did)
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    lang_storage.update_one({"number": number, "name": name, "lang": lang, "status": "WAITING_FOR_TRANSLATION"},
                            {"$set": {"piece number": len(doc.paragraphs)}})
    return ids


def push_to_file_storage(path, file):
    """
    :param path: file path
    :param file: file
    :return: Nothing
    """
    file.save(path)


def push_to_db(number, name, status, lang, importance=0, pieces_count=None, path=None, orig_path=None, file_data=None,
               tags=None,
               freedom=True, index=None, to_lang="RUS", translator=None, piece_begin=None, piece_end=None, txt=None,
               translated_txt=None, translation_status="UNDONE", chief=None):
    """
    :param number: id number of document
    :param name: file name
    :param status: one of TRANSLATED/NEED_CHECK/PIECE/WAITING_PIECE/WAITING_FOR_TRANSLATION
    :param lang: language one of ENG, RUS, ESP, JAP, etc.
    :param importance: number, how this doc is needed
    :param pieces_count: amount of pieces in document
    :param path: file path in the filesystem
    :param orig_path: path to original file
    :param file_data: file
    :param tags: additional tags
    :param freedom: is piece not taken True/False
    :param index: piece index
    :param to_lang: language file is translated to
    :param translator: mongo id of translator
    :param piece_begin: piece beginning paragraph
    :param piece_end: piece ending paragraph
    :param txt: piece text
    :param translated_txt: translated piece
    :param translation_status: whether translation done or not (DONE/UNDONE)
    :param chief: translates verifier
    :return: file mongo id in DB
    """

    file = None
    client = MongoClient()
    db = client.highlight

    if status == "WAITING_FOR_TRANSLATION":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "orig path": orig_path,
                "piece number": pieces_count,
                "tags": tags,
                "importance": importance,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}
        push_to_file_storage(orig_path, file_data)
        split_to_pieces(number, name, lang, file_data)

    elif status == "WAITING_PIECE":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "txt": txt,
                "index": index,
                "freedom": freedom,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}

    elif status == "PIECE":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "piece begin": piece_begin,
                "piece end": piece_end,
                "txt": txt,
                "translated txt": translated_txt,
                "translator": translator,
                "to lang": to_lang,
                "translation status": translation_status,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}

    elif status == "NEED_CHECK" or status == "TRANSLATED":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "path": path,
                "orig path": orig_path,
                "to lang": to_lang,
                "tags": tags,
                "translator": translator,
                "chief": chief,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}
        push_to_file_storage(path, file_data)

    lang_storage = db.files_info
    file_id = lang_storage.insert_one(file).inserted_id

    return file_id


def update_importance(doc_id):
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    lang_storage.update_one({"_id": ObjectId(doc_id)}, {"$inc": {"importance": 1}})


def update_pieces(user_id, doc_id, pieces_ids, to_lang="RUS"):
    """
    :param user_id: mongo id of user
    :param doc_id: mongo id of file
    :param pieces_ids: mongo ids of pieces list
    :param to_lang: language file is translated to
    :return: response code
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    pieces = list()
    document = lang_storage.find_one({"_id": ObjectId(doc_id)})
    no_intersections = True
    for piece_id in pieces_ids:
        p = lang_storage.find_one({"_id": ObjectId(piece_id)})
        pieces.append(p)
        if not p["freedom"]:
            no_intersections = False
            return ("3000",)
    pieces = sorted(pieces, key=lambda a: a["index"])
    txt = [pieces[0]["txt"]]
    for i in range(1, len(pieces)):
        if pieces[i]["index"] - pieces[i - 1]["index"] != 1:
            return ("3001",)
        txt.append(pieces[i]["txt"])
    begin_index = pieces[0]["index"]
    end_index = pieces[-1]["index"]
    if not no_intersections:
        return ("3000",)
    else:
        acc = db.accounts
        new_piece = {
            "name": document["name"],
            "indexes": [range(begin_index, end_index+1)],
            "reservation date": datetime.datetime.utcnow()
        }
        acc.update_one({"_id": ObjectId(user_id)}, {"$push": {"pieces": new_piece}})
        did1 = push_to_db(number=document["number"],
                   name=document["name"],
                   lang=document["lang"],
                   piece_begin=begin_index,
                   piece_end=end_index,
                   txt=txt,
                   translated_txt=None,
                   translator=user_id,
                   to_lang=to_lang,
                   status="PIECE",
                   translation_status="UNDONE")
        for p in pieces:
            lang_storage.update_one({"_id": p["_id"]},
                                    {"$set": {"freedom": False, "lastModified": datetime.datetime.utcnow()}})
        return ("OK", did1)


def update_docs(name, doc, lang, tags):
    """
    :param name: file name
    :param doc: file data .docx
    :param lang: language
    :param tags: tags as array
    :return: file mongo id
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    did = push_to_db(lang_storage.count_documents({"status": "WAITING_FOR_TRANSLATION"}) + 1, name,
                     "WAITING_FOR_TRANSLATION", lang, tags=tags, pieces_count=0, importance=0,
                     orig_path=os.getcwd() + os.path.sep + 'file_storage' + os.path.sep + 'original' + os.path.sep + name,
                     file_data=doc)
    return did


def update_translating_pieces(piece_id, tr_txt=None, tr_stat="UNDONE"):
    """
    :param piece_id: mongo id of piece
    :param tr_txt: translated text string
    :param tr_stat: status of translation DONE/UNDONE
    :return: None or translated doc id
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    lang_storage.update_one({"_id": ObjectId(piece_id)}, {
        "$set": {"translated txt": tr_txt, "translation status": tr_stat, "lastModified": datetime.datetime.utcnow()}})
    if tr_stat == "DONE":
        ps = lang_storage.find_one({"_id": ObjectId(piece_id)})
        acc = db.accounts
        acc.update_one({"_id": ObjectId(ps["translator"])}, {"$inc": {"translated": 1}})
        doc = lang_storage.find_one(
            {"number": ps["number"], "name": ps["name"], "lang": ps["lang"], "status": "WAITING_FOR_TRANSLATION"})
        pieces_count = doc["piece number"]
        taken_pieces_indexes = []
        pss = lang_storage.find({"number": ps["number"], "name": ps["name"], "lang": ps["lang"], "status": "PIECE",
                                 "translation status": "DONE"})
        pss = sorted(pss, key=lambda a: a["piece begin"])
        for p in pss:
            taken_pieces_indexes.extend(range(p["piece begin"], p["piece end"] + 1))
        if pieces_count <= len(taken_pieces_indexes):
            return create_translated_unverified_docs(pss, doc, ps, acc)
        else:
            return None


def create_translated_unverified_docs(pieces, doc, ps, acc):
    """
    :param pieces: all text pieces (translated)
    :param doc: document in db format
    :param ps: one of pieces
    :param acc: accounts db
    :return: mongo id of created element
    """
    file_data = find_file_by_path(doc["orig path"])
    for p in pieces:
        txts = p["translated txt"]
        for i in range(p["piece begin"], p["piece end"] + 1):
            file_data.paragraphs[i].text = txts[i]

    the_stat = "TRANSLATED"

    for tr_id in list({p["translator"] for p in pieces}):
        if acc.find_one({"_id": tr_id})["status"] == "translator":
            the_stat = "NEED_CHECK"

    did = push_to_db(doc["number"], doc["name"], the_stat, doc["lang"], orig_path=doc["orig path"],
                     path=os.getcwd() + os.path.sep + 'file_storage' + os.path.sep + 'translated' + os.path.sep + doc[
                         "name"], to_lang=ps["to lang"], tags=doc["tags"], translator=list({p["translator"] for p in pieces}),
                     file_data=file_data)
    for p in pieces:
        delete_from_db(p["_id"])
    return did


def delete_from_doc_storage(path):
    os.remove(path)


def delete_from_db(doc_id):
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    doc = lang_storage.find_one({"_id": doc_id})
    if "path" in doc.keys():
        delete_from_doc_storage(doc["path"])
    lang_storage.delete_one({"_id": doc_id})


def find_file_by_path(path):
    document = Document(path)
    return document


def find_pieces(user_id):
    """
    :param user_id: user mongo id
    :return: all pieces taken by user
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    pieces = sorted(lang_storage.find({"translator": user_id, "status": "PIECE", "translation status": "UNDONE"}),
                    key=lambda a: a["lastModified"], reverse=True)
    return pieces


def find_doc_by_lang(lang):
    """
    :param lang: language
    :return: array of tuples where first element is document name, second is list of pieces for this document
    :description: finds pieces as docs by language
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    docs = dict()
    docs_o = dict()
    for piece in lang_storage.find({"lang": lang, "status": "WAITING_PIECE"}):
        orig_doc = lang_storage.find_one({"name": piece["name"], "number": piece["number"], "lang": piece["lang"], "status": "WAITING_FOR_TRANSLATION"})
        if piece["name"] in docs.keys():
            docs[piece["name"]].append(piece)
        else:
            docs[piece["name"]] = [piece]
            docs_o[piece["name"]] = [orig_doc]

    out = list()
    for key in docs.keys():
        docs[key] = sorted(docs[key], key=lambda a: a["index"])
        out.append([key, docs[key], docs_o])
    sorted(out, key=lambda a: a[2], reverse=True)

    return out


def find_complete_docs_by_lang(lang):
    """
    :param lang: language
    :return: array of docs (id, name, tags, progress)
    :description: finds docs by language
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    docs = list()
    for doc in lang_storage.find({"lang": lang, "status": { "$in": ["WAITING_FOR_TRANSLATION", "NEED_CHECK", "TRANSLATED"]}}):
        if doc["status"] in {"NEED_CHECK", "TRANSLATED"}:
            docs.append([doc["_id"], doc["name"], doc["tags"], 1])
        else:
            pieces_count = doc["piece number"]
            taken_pieces_indexes = []
            pss = lang_storage.find({"number": doc["number"], "name": doc["name"], "lang": doc["lang"], "status": "PIECE",
                                     "translation status": "DONE"})
            pss = sorted(pss, key=lambda a: a["piece begin"])
            for p in pss:
                taken_pieces_indexes.extend(range(p["piece begin"], p["piece end"] + 1))
            docs.append([doc["_id"], doc["name"], doc["tags"], len(taken_pieces_indexes) / pieces_count])

    return docs


def get_from_db(search, tags, status=None):
    """
    :param search: user input
    :param tags: tags
    :param status: list of statuses
    :return: matching docs of id, name, tags, status, path, etc.
    """
    if status is None:
        status = {"TRANSLATED", "NEED_CHECK", "WAITING_FOR_TRANSLATION"}
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    search_set = set(search)
    hl = []
    for i in search.split(" "):
        hl.extend(i.split("_"))
    search_set_words = set(hl)
    search_tags = set(tags.split(","))
    matching_docs = list()
    for doc in lang_storage.find({"status": {"$in": ["TRANSLATED", "NEED_CHECK", "WAITING_FOR_TRANSLATION"]}}):
        doc_stat = doc["status"]
        relev = 0
        if doc_stat == "TRANSLATED":
            relev += 100
        elif doc_stat == "NEED_CHECK":
            relev += 50
        elif doc_stat == "WAITING_FOR_TRANSLATION":
            relev += 0
        doc_name_set = set(doc["name"])
        doc_tags_set = set(doc["tags"])
        doc_id = doc["number"]
        tag_inrsc = search_tags.intersection(doc_tags_set)
        name_inrsc = search_set.intersection(doc_name_set)

        if len(search_tags) > 0:
            if len(tag_inrsc) >= len(search_tags) * 0.8:
                relev += len(tag_inrsc) / len(search_tags) * 7

            if doc["lang"] in search_tags:
                relev += 4

        if len(search_set) > 0:
            if len(name_inrsc) >= 0.6 * len(doc_name_set):
                relev += len(name_inrsc) / len(doc_name_set) * 10

            if doc_id in search_set_words:
                relev += 8

        if relev > 0:
            matching_docs.append((relev, doc))

    return list(d for n, d in sorted(matching_docs, key=lambda t: t[0], reverse=True) if d["status"] in status)


def get_for_chief_from_db(search, tags):
    return get_from_db(search, tags, status={"NEED_CHECK"})


def get_users():
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    return acc.find({"verified": False})


def get_docs_and_trans():
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    l_s = db.files_info
    return l_s.count_documents({"status": "WAITING_FOR_TRANSLATION"}), acc.count_documents({"status": {"$in": ["translator", "both"]}, "verified": True}), l_s.count_documents({"status": "TRANSLATED"})


def get_translators_stat():
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    return [t for t in acc.find({"status": {"$in": ["translator", "both"]}, "verified": True})]


def get_file_stat():
    client = MongoClient()
    db = client.highlight
    l_s = db.files_info
    docs = []
    for t in l_s.find({"status": {"$in": ["TRANSLATED", "NEED_CHECK", "WAITING_FOR_TRANSLATION"]}}):
        if t["status"] in {"TRANSLATED", "NEED_CHECK"}:
            docs.append([t["name"], t["status"], l_s.find_one({"name": t["name"], "number": t["number"], "lang": t["lang"], "status": "WAITING_FOR_TRANSLATION"})["importance"]])
        else:
            docs.append([t["name"], [l_s.count_documents({"name": t["name"], "number": t["number"], "lang": t["lang"], "status": "PIECE", "translation status": "DONE"}), t["piece number"]], t["importance"]])
    return docs


def test():
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    # acc.create_index([('login', pm.ASCENDING)], unique=True)
    # did = register("seva", "obvious", "obvious", "obvious", "obvious", "seva", "tester", "obvious")
    # print(did)
    # did1 = get_users()[0]
    # print(did1["_id"])
    # verify(did1["_id"], "NICE")
    # print(log_in("seva", "tester"))
    # for i in acc.find():
    #     pprint.pprint(i)
    # acc.delete_one({"_id": log_in("seva", "tester")})


if __name__ == '__main__':
    test()
