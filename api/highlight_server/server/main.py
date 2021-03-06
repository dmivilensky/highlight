import os
import pprint
from bson.objectid import ObjectId

from docx import Document
from pymongo import MongoClient
import pymongo as pm
import datetime
import django
import random
import shutil as sh
import pdfminer.high_level as phl
import re

if __name__ == "main":
    from logger import Logger
    from filemanager import MergeStatus, FileManager
else:
    from .logger import Logger
    from .filemanager import MergeStatus, FileManager

PATH_TO_FILES = os.path.dirname(os.path.realpath(__file__)) + "/../../../files"

BOOL_TO_ABB = ["ENG", "GER", "FRE", "ESP", "ITA", "JAP", "CHI"]
FM = FileManager(PATH_TO_FILES)


def verify_file(doc_id, user_id, file_path=None):
    """
    :param doc_id: document mongo id
    :param user_id: user mongo id
    :param file_path: path to save file returned by php script
    :param file_status: status of the file after verification
    :return: code
    :structure: dict('code': string)
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    acc = db.accounts
    user = acc.find_one({"_id": ObjectId(user_id)})
    if is_there_any_body(user_id):
        if user["status"] in {"chief", "both"}:
            if not(file_path is None):
                file = lang_storage.find_one({"_id": ObjectId(doc_id)})
                delete_from_doc_storage(file["path"])
                # push_to_file_storage(file["path"], file_data)
                # f = open('program_logs.txt', 'w')
                # f.write(file["path"])
                # f.close()
                lang_storage.update_one({"_id": ObjectId(doc_id)},
                                        {"$set": {"status": "TRANSLATED", "chief": user_id, "path": file_path}})
            else:
                lang_storage.update_one({"_id": ObjectId(doc_id)}, {"$set": {"status": "TRANSLATED", "chief": user_id}})
            return {"code": "OK"}
        elif user["status"] == "verif":
            if not (file_path is None):
                file = lang_storage.find_one({"_id": ObjectId(doc_id)})
                delete_from_doc_storage(file["path"])
                # push_to_file_storage(file["path"], file_data)
                # f = open('program_logs.txt', 'w')
                # f.write(file["path"])
                # f.close()
                lang_storage.update_one({"_id": ObjectId(doc_id)},
                                        {"$set": {"substatus": "DONE", "path": file_path}})
            else:
                lang_storage.update_one({"_id": ObjectId(doc_id)},
                                        {"$set": {"substatus": "DONE"}})
            return {"code": "OK"}
        else:
            return {"code": "2004"}
    else:
        return {"code": "2003"}


def is_there_any_body(uid):
    """
    :param uid: user mongo id
    :return: is user in db
    :structure: bool
    """
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    return not(acc.find_one({"_id": ObjectId(uid), "verified": True}) is None)


def split_to_pieces(number, name, lang, doc_path):
    """
       :param number: id number of document
       :param name: file name
       :param lang: language one of ENG, RUS, ESP, JAP, etc.
       :param doc_path: document path
       :return: pieces ids
       :structure: list(ObjectId Bson)
   """

    if len(doc_path.split("/")) and doc_path.split("/")[-2] != "files":
        integrated = True
    else:
        integrated = False
    ids = list()
    counter = 0
    lgr = Logger()
    lgr.log("log", "splitting", "path " + (doc_path.split("/")[-1] if not integrated else doc_path.split("/")[-2] + "/" + doc_path.split("/")[-1]))
    piece_pages = FM.split_pdf((doc_path.split("/")[-1] if not integrated else doc_path.split("/")[-2] + "/" + doc_path.split("/")[-1]))
    lgr.log("log", "splitting", "entry point")
    for i in range(len(piece_pages)):
        lgr.log("log", "splitting", "try " + str(i))
        did = push_to_db(number, name, "WAITING_PIECE", lang, txt_path=piece_pages[i], index=counter, freedom=True)
        lgr.log("log", "splitting", "try " + str(i))
        ids.append(did)
        counter += 1
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    lgr.log("log", "updating doc with pcount", "num " + str(number) + "  name " + name + "  lang " + lang + "  count " + str(counter))
    udd = lang_storage.update_one({"number": number, "name": name, "lang": lang, "status": "WAITING_FOR_TRANSLATION"},{"$set": {"piece_number": counter}})
    lgr.log("log", "updating doc with pcount", "upd_d " + str([udd.modified_count, udd.matched_count, udd.upserted_id, udd.acknowledged]))
    return ids, piece_pages


def push_to_file_storage(path, file):
    """
    :param path: file path
    :param file: file
    :return: Nothing
    """
    file.save(path)
    lgr = Logger()
    lgr.log("log", "loader status: ", "file saved")


def extract_pdf2text(fpath, db_file, lang_storage, ptxts, orig=True):
    db_file = lang_storage.find_one({"_id": db_file["_id"]})
    db_file = pdf2text_core(fpath, db_file, ptxts, orig)

    lang_storage.replace_one({"number": db_file["number"], "name": db_file["name"], "lang": db_file["lang"], "status": db_file["status"]},
                            db_file)
    return {"code": "OK"}


def pdf2txt_from_doc(doc, ptxts):
    return pdf2text_core((doc["orig_path"] if doc["status"] == "WAITING_FOR_TRANSLATION" else doc["path"]), doc, ptxts, orig=(True if doc["status"] == "WAITING_FOR_TRANSLATION" else False))


def pdf2text_core(fpath, db_file, ptxts, orig):
    text_preview = ""
    lgr = Logger()
    lgr.log("log", "loader status: ", "extracting text from pdf")
    piece_pages = ptxts
    all_text = []
    has_preview = False
    for i in range(len(piece_pages)):
        text = phl.extract_text(PATH_TO_FILES + "/" + piece_pages[i]).strip()
        if not (has_preview):
            has_preview = True
            try:
                text_preview = list(filter(lambda p: p != "" and len(p.split(" ")) >= 50,
                                           map(lambda p: p.strip(), text.split("\n\n"))))[0].strip()
            except IndexError:
                try:
                    text_preview = \
                        list(filter(lambda p: p != "" and len(p.split(" ")) >= 10,
                                    map(lambda p: p.strip(), text.split("\n\n"))))[
                            0].strip()
                except IndexError:
                    pass
        all_text.append(text)
    path = PATH_TO_FILES + "/" + "pdf2text" + ("orig" if orig else "trans") + str(random.randint(0, 99999)) + str(
        random.randint(0, 99999)) + ".txt"
    while os.path.isfile(path):
        path = path[:-len(path.split(".")[-1]) + 1] + str(random.randint(0, 99999)) + ".txt"
    lgr = Logger()
    lgr.log("log", "loader status: ", "saving text")
    with open(path, 'w+') as destination:
        destination.write("\n//||\\\\\n".join(all_text))
    if orig:
        db_file["orig_txt_path"] = path
        db_file["orig_preview"] = "...\n" + text_preview + "\n..."
    else:
        db_file["txt_path"] = path
        db_file["preview"] = "...\n" + text_preview + "\n..."
    return db_file


def indexing(db, file, file_id, orig=True):
    lang_storage = MongoClient().highlight.files_info
    file = lang_storage.find_one({"_id": file_id})
    index = db.index_holder
    index_file = index.find_one({"name": "index"})
    with open((file["orig_txt_path"] if orig else file["txt_path"]), 'r') as destination:
        words = set()
        for line in destination:
            if line.strip() != "":
                words |= set(filter(lambda p: p != "" and p != "//||\\\\", map(lambda p: p.strip().lower(), re.sub("[^\w]", " ", line).split())))
        if index_file is None:
            index_file = {
                "name": "index"
            }
            w_in_a = dict()
            for word in words:
                w_in_a[word] = [file_id]
            index_file["word2articles"] = w_in_a
            index_file["word_count"] = len(w_in_a)
            index_id = index.insert_one(index_file).inserted_id
        else:
            w_in_a = index_file["word2articles"]
            for word in words:
                if word in w_in_a.keys():
                    w_in_a[word].append(file_id)
                else:
                    w_in_a[word] = [file_id]
            index_id = index.update_one({"name": "index"},
                                    {"$set": {"word2articles": w_in_a, "word_count": len(w_in_a)}}).upserted_id
    return index_id


def combine_indexing_for_update(doc, is_splitted=True, is_extracted=True):
    """
    :param dict doc: document in mingodb format
    :param bool is_splitted: whether file is splitted to pages
    :param bool is_extracted: whether file is extracted to texts
    :return: dict() - mongodb file
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    if is_splitted:
        pss = lang_storage.find({"number": doc["number"], "name": doc["name"], "lang": doc["lang"], "status": "WAITING_PIECE"})
        pids, ptxts = zip(*[(p["_id"], p["txt_path"]) for p in pss])
    else:
        pids, ptxts = split_to_pieces(doc["number"], doc["name"], doc["lang"], doc["orig_path"])
        doc = lang_storage.find_one({"_id": doc["_id"]})

    if is_extracted:
        pass
    else:
        doc = pdf2txt_from_doc(doc, ptxts)
    indexing(db, doc, doc["_id"], orig=(True if doc["status"] == "WAITING_FOR_TRANSLATION" else False))
    return doc


def push_to_db(number, name, status, lang, importance=0, pieces_count=None, path=None, orig_path=None, file_data=None,
               tags=None,
               freedom=True, index=None, to_lang="RUS", translator=None, piece_begin=None, piece_end=None, txt_path=None,
               translated_txt_path=None, translation_status="UNDONE", chief=None, author=None, journal=None, abstract=None, journal_link=None, orig_txt_path=None, orig_preview=None):
    """
    :param int number: id number of document
    :param str name: file name
    :param str status: one of TRANSLATED/NEED_CHECK/PIECE/WAITING_PIECE/WAITING_FOR_TRANSLATION/MARKUP
    :param str lang: language one of ENG, RUS, ESP, JAP, etc.
    :param int importance: number, how this doc is needed
    :param int pieces_count: amount of pieces in document
    :param str path: file path in the filesystem
    :param str orig_path: path to original file
    :param file_data: file
    :param str tags: additional tags
    :param freedom: is piece not taken True/False
    :param index: piece index
    :param to_lang: language file is translated to
    :param translator: mongo id of translator
    :param piece_begin: piece beginning paragraph
    :param piece_end: piece ending paragraph
    :param txt_path: piece text
    :param translated_txt_path: translated piece
    :param translation_status: whether translation done or not (DONE/UNDONE)
    :param chief: translates verifier
    :param author: article author
    :param journal: journal in which article was published
    :param abstract: abstract
    :param journal_link: link to coresponding journal
    :param orig_preview: preview txt from WAITING_FOR_TRANSLATION
    :param orig_txt_path: txt file path from WAITING_FOR_TRANSLATION
    :return: file mongo id in DB
    :structure: ObjectId Bson
    """

    file = None
    client = MongoClient()
    db = client.highlight

    if status == "WAITING_FOR_TRANSLATION":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "orig_path": orig_path,
                "piece_number": pieces_count,
                "tags": tags,
                "importance": importance,
                "status": status,
                "author": author,
                "abstract": abstract,
                "journal": journal,
                "journal_link": journal_link,
                "orig_txt_path": "",
                "orig_preview": "",
                "lastModified": datetime.datetime.utcnow()}
        lgr = Logger()
        lgr.log("log", "loader status: ", "saved to db")

    elif status == "WAITING_PIECE":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "txt_path": txt_path,
                "index": index,
                "freedom": freedom,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}

    elif status == "PIECE":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "piece_begin": piece_begin,
                "piece_end": piece_end,
                "txt_path": txt_path,
                "translated_txt_path": translated_txt_path,
                "translator": translator,
                "to_lang": to_lang,
                "translation_status": translation_status,
                "status": status,
                "lastModified": datetime.datetime.utcnow()}

    elif status == "NEED_CHECK" or status == "TRANSLATED":
        file = {"number": number,
                "name": name,
                "lang": lang,
                "path": path,
                "orig_path": orig_path,
                "to_lang": to_lang,
                "tags": tags,
                "importance": importance,
                "translator": translator,
                "chief": chief,
                "status": status,
                "substatus": "MARKUP",
                "author": author,
                "abstract": abstract,
                "journal": journal,
                "journal_link": journal_link,
                "orig_txt_path": orig_txt_path,
                "orig_preview": orig_preview,
                "txt_path": "",
                "preview": "",
                "lastModified": datetime.datetime.utcnow()}

    lang_storage = db.files_info
    file_id = lang_storage.insert_one(file).inserted_id
    client.close()

    lgr = Logger()
    lgr.log("log", "update db", "additional operations")

    if status == "WAITING_FOR_TRANSLATION":
        try:
            pids, ptxts = split_to_pieces(number, name, lang, orig_path)
        except Exception as e:
            print(e)
            lgr.log("log", "update db splitting", str(e))

    if status in {"WAITING_FOR_TRANSLATION", "NEED_CHECK", "TRANSLATED"}:
        try:
            if status in {"TRANSLATED", "NEED_CHECK"}:
                extract_pdf2text(path, file, lang_storage, [path.split("/")[-1] if len(path.split("/")) < 2 or path.split("/")[-2] == "files" else ("/".join(path.split("/")[(path.split("/").index("files")+1):]))], orig=False)
            else:
                extract_pdf2text(orig_path, file, lang_storage, ptxts)
            indexing(db, file, file_id, orig=(True if status == "WAITING_FOR_TRANSLATION" else False))
        except FileNotFoundError as e:
            print(e)
            lgr.log("log", "update db indexing", str(e))
        except IsADirectoryError as e:
            print(e)
            lgr.log("log", "update db indexing", str(e))
        except UnboundLocalError as e:
            print(e)
            lgr.log("log", "update db indexing", str(e))
        except Exception as e:
            print(e)
            lgr.log("log", "update db indexing", str(e))

    return file_id


def update_importance(doc_id):
    """
    :param doc_id: document mongo id
    :return: code
    :structure: dict('code': string)
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    lang_storage.update_one({"_id": ObjectId(doc_id)}, {"$inc": {"importance": 1}})
    return {"code": "OK"}


def update_pieces(user_id, doc_id, pieces_ids, to_lang="RUS"):
    """
    :param user_id: mongo id of user
    :param doc_id: mongo id of file
    :param pieces_ids: mongo ids of pieces list
    :param to_lang: language file is translated to
    :return: response code
    :structure: dict('code': string)
    :description: creates piece taken by user
    """
    lgr = Logger()
    lgr.log("log", "update pieces", "entry main")
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    pieces = list()
    document = lang_storage.find_one({"_id": ObjectId(doc_id)})
    no_intersections = True
    pieces_ids = list(filter(lambda pid: pid != "not", pieces_ids))
    for piece_id in pieces_ids:
        p = lang_storage.find_one({"_id": ObjectId(piece_id)})
        pieces.append(p)
        if not p["freedom"]:
            no_intersections = False
            return {"code": "3000"}
    lgr.log("log", "update pieces", "intersection checked")
    pieces = sorted(pieces, key=lambda a: a["index"])
    txt = [pieces[0]["txt_path"]]
    for i in range(1, len(pieces)):
        if pieces[i]["index"] - pieces[i - 1]["index"] != 1:
            return {"code": "3001"}
        txt.append(pieces[i]["txt_path"])
    begin_index = pieces[0]["index"]
    end_index = pieces[-1]["index"]
    if not no_intersections:
        return {"code": "3000"}
    else:
        lgr.log("log", "update pieces", "before compose, files " + ",".join(txt))
        txt_real = FM.compose_files(txt, status=MergeStatus.piece, delete=False)
        lgr.log("log", "update pieces", "after compose")
        tname = FM.create_path("translated")
        lgr.log("log", "update pieces", "lets copy " + PATH_TO_FILES + "/" + txt_real + " to " + PATH_TO_FILES + "/" + tname)
        sh.copy(PATH_TO_FILES + "/" + txt_real, PATH_TO_FILES + "/" + tname)
        acc = db.accounts
        new_piece = {
            "name": document["name"],
            "indexes": list(range(begin_index, end_index+1)),
            "reservation_date": datetime.datetime.utcnow()
        }
        acc.update_one({"_id": ObjectId(user_id)}, {"$push": {"pieces": new_piece}})
        did1 = push_to_db(number=document["number"],
                          name=document["name"],
                          lang=document["lang"],
                          piece_begin=begin_index,
                          piece_end=end_index,
                          txt_path=txt_real,
                          translated_txt_path=tname,
                          translator=user_id,
                          to_lang=to_lang,
                          status="PIECE",
                          translation_status="UNDONE")
        for p in pieces:
            lang_storage.update_one({"_id": p["_id"]},
                                    {"$set": {"freedom": False, "lastModified": datetime.datetime.utcnow()}})
        return {"id": str(did1), "code": "OK"}


def update_docs(name, doc, lang, tags, author=None, abstract=None, journal=None, jl=None, path=None):
    """
    :param name: file name
    :param doc: file data .pdf
    :param lang: language
    :param tags: tags as array
    :param author: author
    :param abstract: abstract
    :param journal: journal
    :param path: path to file
    :param jl: link to journal
    :return: file mongo id
    :structure: dict('code': string, 'id': string)
    :description: pushes loaded by admin file to db
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    pf = lang_storage.find_one({"name": name, "status": "WAITING_FOR_TRANSLATION", "orig_path": path})
    if not(pf is None):
        return {"code": "7777"}
    lgr = Logger()
    lgr.log("log", "loader status: ", "saving to db")
    did = push_to_db(lang_storage.count_documents({"status": "WAITING_FOR_TRANSLATION"}) + 1, name,
                     "WAITING_FOR_TRANSLATION", lang, tags=tags, pieces_count=0, importance=0,
                     orig_path=path, file_data=doc, abstract=abstract, author=author, journal=journal, journal_link=jl)
    return {"id": str(did), "code": "OK"}


def update_translating_pieces(piece_id, tr_txt=None, tr_stat="UNDONE"):
    """
    :param piece_id: mongo id of piece
    :param tr_txt: translated text string
    :param tr_stat: status of translation DONE/UNDONE
    :return: error or translated doc id
    :structure: dict('code': string, 'id': string)
    :description: updates users piece information: whether translation done or not and ads translated text to piece. If all pieces of certain document are translated - creates translated document and deletes pieces
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    if tr_txt is None:
        lang_storage.update_one({"_id": ObjectId(piece_id)}, {
            "$set": {"translation_status": tr_stat, "lastModified": datetime.datetime.utcnow()}})
    else:
        lang_storage.update_one({"_id": ObjectId(piece_id)}, {
            "$set": {"translated_txt_path": tr_txt, "translation_status": tr_stat, "lastModified": datetime.datetime.utcnow()}})
    if tr_stat == "DONE":
        lgr = Logger()
        lgr.log("log", "final file", "entry")
        ps = lang_storage.find_one({"_id": ObjectId(piece_id)})
        acc = db.accounts
        acc.update_one({"_id": ObjectId(ps["translator"])}, {"$inc": {"translated": 1}})
        doc = lang_storage.find_one(
            {"number": ps["number"], "name": ps["name"], "lang": ps["lang"], "status": "WAITING_FOR_TRANSLATION"})
        pieces_count = doc["piece_number"]
        taken_pieces_indexes = []
        pss = list(lang_storage.find({"number": ps["number"], "name": ps["name"], "lang": ps["lang"], "status": "PIECE",
                                 "translation_status": "DONE"}))
        lgr.log("log", "final file", str(pss))
        pss = sorted(pss, key=lambda a: a["piece_begin"])
        for p in pss:
            taken_pieces_indexes.extend(range(p["piece_begin"], p["piece_end"] + 1))
        lgr.log("log", "final file", str(pieces_count) + " " + str(taken_pieces_indexes))
        # return {"pc": pieces_count, "tpi": len(taken_pieces_indexes)}
        if pieces_count <= len(taken_pieces_indexes):
            return {"id": str(create_translated_unverified_docs(pss, doc, ps, acc, lang_storage=lang_storage)), "code": "OK"}
        else:
            return {"code": "OK", "document": "3002"}
    else:
        return {"code": "OK", "document": tr_txt}


def create_translated_unverified_docs(pieces, doc, ps, acc, lang_storage=None):
    """
    :param pieces: all text pieces (translated)
    :param doc: document in db format
    :param ps: one of pieces
    :param acc: accounts db
    :return: mongo id of created element
    :description: creates translated document from pieces
    """

    lgr = Logger()
    lgr.log("log", "compose", "entry")
    lgr.log("log", "compose", "files " + "/".join([p["translated_txt_path"] for p in pieces]))
    file_path = FM.compose_files([p["translated_txt_path"] for p in pieces], status=MergeStatus.composition, delete=False)
    lgr.log("log", "compose", "result" + str(file_path))

    the_stat = "TRANSLATED"

    chief_id = []
    for tr_id in list({p["translator"] for p in pieces}):
        cOt = acc.find_one({"_id": ObjectId(tr_id)})
        if cOt["status"] == "translator":
            the_stat = "NEED_CHECK"
        else:
            chief_id.append(tr_id)

    lgr.log("log", "compose", "ready to push")
    lgr.log("log", "compose", "data: " + str(doc["number"]) + " " + doc["name"])
    did = push_to_db(doc["number"], doc["name"], the_stat, doc["lang"], orig_path=doc["orig_path"],
                     path=PATH_TO_FILES + "/" + file_path, to_lang=ps["to_lang"], tags=doc["tags"], importance=doc["importance"], translator=list({p["translator"] for p in pieces}), chief=chief_id, author=doc["author"], abstract=doc["abstract"], journal=doc["journal"], journal_link=doc["journal_link"], orig_preview=doc["orig_preview"], orig_txt_path=doc["orig_txt_path"])
    lgr.log("log", "compose", "done push")
    # for p in pieces:
    #     delete_from_db(p["_id"])
    # if not(lang_storage is None):
    #     for p in list(lang_storage.find({"number": doc["number"], "name": doc["name"], "lang": doc["lang"], "status": "WAITING_PIECE"})):
    #         delete_from_db(p["_id"])
    delete_from_db(doc["_id"], with_path=False)
    return did


def delete_from_db_all(did):
    """
    :param did: document id
    :return: code
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    doc = lang_storage.find_one({"_id": ObjectId(did)})
    ll = []
    for p in list(lang_storage.find(
            {"number": doc["number"], "name": doc["name"], "lang": doc["lang"]
            # ,
            # "status": {"$in": ["WAITING_PIECE", "PIECE", "NEED_CHECK", "TRANSLATED", "MARKUP"]}
            })):
        ll.append(str(p["_id"]))
        delete_from_db(p["_id"])
    return {"code": ll}


def delete_from_doc_storage(path):
    """
    :param path: file path
    :return: code
    """
    if os.path.isfile(path):
        os.remove(path)
    elif os.path.isfile(PATH_TO_FILES + "/" + path):
        os.remove(PATH_TO_FILES + "/" + path)
    return {"code": "OK"}


def delete_from_db(doc_id, with_path=True, recursively=False):
    """
    :param doc_id: file mongo id
    :param with_path: delete file in storage as well
    :return: Nothing
    :description: deletes file from db and storage if any
    """
    client = MongoClient()
    db = client.highlight
    lang_storage = db.files_info
    doc = lang_storage.find_one({"_id": doc_id})

    if recursively:
        ll = []
        for p in list(lang_storage.find(
            {"number": doc["number"], "name": doc["name"], "lang": doc["lang"]
            # ,
            # "status": {"$in": ["WAITING_PIECE", "PIECE", "NEED_CHECK", "TRANSLATED", "MARKUP"]}
            })):
            ll.append(str(p["_id"]))
            delete_from_db(p["_id"], with_path=with_path)

    if with_path:
        if "path" in doc.keys():
            delete_from_doc_storage(doc["path"])
        if "orig_path" in doc.keys():
            delete_from_doc_storage(doc["orig_path"])
        if "txt_path" in doc.keys():
            delete_from_doc_storage(doc["txt_path"])
        if "translated_txt_path" in doc.keys():
            delete_from_doc_storage(doc["translated_txt_path"])
        if "orig_txt_path" in doc.keys():
            delete_from_doc_storage(doc["orig_txt_path"])
    lang_storage.delete_one({"_id": doc_id})


def find_file_by_path(path):
    """
    :param path: file path
    :return: .docx document
    """
    return Document(path)


def test():
    client = MongoClient()
    db = client.highlight
    acc = db.accounts
    #ids = acc.insert_one({"name": "seva", "login": "lol", "password": "kek", "verified": True})
    #print(ids)
    #print(list(acc.find()))
    d = Document("files/new_file4283.docx")
    for p in d.paragraphs:
        print(p.text)


if __name__ == '__main__':
    test()
