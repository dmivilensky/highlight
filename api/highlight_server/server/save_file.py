from django.http import HttpResponse
from django.views.decorators.csrf import csrf_exempt

from .forms import UploadFileForm
from .utils import get_params, handle_uploaded_file
import docx


@csrf_exempt
def save_file(request):
    if request.method == "POST":
            form = UploadFileForm(get_params(request), request.FILES)
            if form.is_valid():
                path = handle_uploaded_file(request.FILES['file'])
            else:
                path = ""
    return path


def test():
    d = docx.Document()
    d.add_paragraph('Lorem ipsum dolor sit amet.')
    d.paragraphs[0].text = "fvdfvfdvfdd\ngvgerveve"
    for i in range(len(d.paragraphs)):
        print(str(i) + ": " + d.paragraphs[i].text)
    print(i+1)


if __name__ == '__main__':
    test()